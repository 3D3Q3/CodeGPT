#!/usr/bin/env python3

import sys
import os
from datetime import datetime
from pathlib import Path
import pypdf
from docx import Document
import mammoth

def get_unique_output_name(extension):
    """Generate a unique output filename based on timestamp."""
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"merged_docs_{timestamp}{extension}"

def merge_pdfs(pdf_files):
    """Merge PDF files with a blank page separator."""
    merger = pypdf.PdfMerger()
    
    # Create a blank page for separation
    blank_pdf = pypdf.PdfWriter()
    blank_pdf.add_blank_page(width=612, height=792)  # Standard letter size
    blank_path = "temp_blank.pdf"
    with open(blank_path, "wb") as temp_file:
        blank_pdf.write(temp_file)
    
    # Merge PDFs with separator
    for pdf_file in pdf_files:
        merger.append(pdf_file)
        merger.append(blank_path)  # Add separator
    
    output_name = get_unique_output_name(".pdf")
    merger.write(output_name)
    merger.close()
    
    # Clean up temporary file
    os.remove(blank_path)
    return output_name

def convert_docx_to_text(docx_path):
    """Extract text from a DOCX file."""
    doc = Document(docx_path)
    return "\n\n".join([paragraph.text for paragraph in doc.paragraphs])

def merge_to_docx(input_files):
    """Merge various documents into a single DOCX with separators."""
    doc = Document()
    
    for i, file_path in enumerate(input_files):
        if i > 0:  # Add separator before each file except the first
            doc.add_paragraph("=" * 80)  # Horizontal rule
            doc.add_page_break()
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            # Extract text from PDF
            pdf_reader = pypdf.PdfReader(file_path)
            text = "\n\n".join(page.extract_text() for page in pdf_reader.pages)
            doc.add_paragraph(text)
            
        elif file_ext == '.docx':
            # Copy content from DOCX
            src_doc = Document(file_path)
            for element in src_doc.element.body:
                doc.element.body.append(element)
                
        elif file_ext == '.txt':
            # Add text file content
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                doc.add_paragraph(txt_file.read())
    
    output_name = get_unique_output_name(".docx")
    doc.save(output_name)
    return output_name

def main():
    if len(sys.argv) < 2:
        print("Usage: python merge_docs.py file1 [file2 ...]")
        sys.exit(1)
    
    input_files = sys.argv[1:]
    
    # Check if all files exist
    for file_path in input_files:
        if not os.path.exists(file_path):
            print(f"Error: File not found - {file_path}")
            sys.exit(1)
    
    # Check if all files are PDFs for PDF-specific merge
    all_pdfs = all(Path(f).suffix.lower() == '.pdf' for f in input_files)
    
    try:
        if all_pdfs:
            output_file = merge_pdfs(input_files)
        else:
            output_file = merge_to_docx(input_files)
        
        print(f"Successfully merged files into: {output_file}")
        
    except Exception as e:
        print(f"Error during merge: {str(e)}")
        sys.exit(1)

if __name__ == "__main__":
    main()